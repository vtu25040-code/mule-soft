import os
import docx

class ManuscriptExporter:
    """
    Multi-format export engine for generated IEEE research manuscripts.
    Supports DOCX, LaTeX (.tex), Markdown (.md), and PDF.
    """

    @staticmethod
    def export_markdown(manuscript_text: str, output_path: str) -> str:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(manuscript_text)
        return output_path

    @staticmethod
    def export_latex(manuscript_data: dict, output_path: str) -> str:
        title = manuscript_data.get("title", "IEEE Research Paper")
        abstract = manuscript_data.get("abstract", "")
        keywords = manuscript_data.get("keywords", "")
        sections = manuscript_data.get("sections", [])
        references = manuscript_data.get("references", [])
        
        latex_lines = [
            "\\documentclass[conference]{IEEEtran}",
            "\\usepackage{cite}",
            "\\usepackage{amsmath,amssymb,amsfonts}",
            "\\usepackage{graphicx}",
            "\\usepackage{textcomp}",
            "\\usepackage{xcolor}",
            "\\begin{document}",
            f"\\title{{{title}}}",
            "\\author{\\IEEEauthorblockN{Author Name} \\IEEEauthorblockA{\\textit{Department of Computer Science} \\\\ \\textit{Research Institute}\\\\ City, Country}}",
            "\\maketitle",
            "\\begin{abstract}",
            abstract,
            "\\end{abstract}",
            "\\begin{IEEEkeywords}",
            keywords,
            "\\end{IEEEkeywords}"
        ]
        
        for sec in sections:
            latex_lines.append(f"\\section{{{sec['title']}}}")
            clean_content = sec['content'].replace("#", "\\#").replace("_", "\\_").replace("&", "\\&")
            latex_lines.append(clean_content)
            
        latex_lines.append("\\begin{thebibliography}{00}")
        for ref in references:
            latex_lines.append(f"\\bibitem{{b{ref['id']}}} {ref['formatted_citation']}")
        latex_lines.append("\\end{thebibliography}")
        latex_lines.append("\\end{document}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n\n".join(latex_lines))
            
        return output_path

    @staticmethod
    def export_docx(manuscript_data: dict, output_path: str) -> str:
        doc = docx.Document()
        
        title = doc.add_heading(manuscript_data.get("title", "IEEE Research Paper"), level=0)
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("Author Name 1, Author Name 2\nDepartment of Computer Science & Engineering\n").alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Abstract
        p_abs = doc.add_paragraph()
        r_abs_label = p_abs.add_run("Abstract—")
        r_abs_label.bold = True
        p_abs.add_run(manuscript_data.get("abstract", ""))
        
        # Keywords
        p_kw = doc.add_paragraph()
        r_kw_label = p_kw.add_run("Index Terms—")
        r_kw_label.bold = True
        p_kw.add_run(manuscript_data.get("keywords", ""))
        
        for sec in manuscript_data.get("sections", []):
            doc.add_heading(f"{sec['number']}. {sec['title']}", level=1)
            doc.add_paragraph(sec['content'])
            
        doc.add_heading("REFERENCES", level=1)
        for ref in manuscript_data.get("references", []):
            doc.add_paragraph(f"[{ref['id']}] {ref['formatted_citation']}")
            
        doc.save(output_path)
        return output_path
